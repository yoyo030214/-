from docx import Document
import os

def analyze_document(file_path):
    try:
        doc = Document(file_path)
        
        # 输出基本信息
        print("文档基本信息：")
        print(f"文件路径：{file_path}")
        print(f"段落数量：{len(doc.paragraphs)}")
        print("\n")
        
        # 输出文本内容
        print("文档内容：")
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        print("\n".join(content))
        print("\n")
        
        # 分块处理
        chunks = split_into_chunks("\n".join(content), 1000)
        print("文档分块（每块约1000字）：")
        for chunk in chunks:
            print("-------------------")
            print(chunk)
            print("-------------------\n")
            
    except Exception as e:
        print(f"处理文档时出错：{str(e)}")

def split_into_chunks(text, chunk_size):
    chunks = []
    sentences = text.split("。")
    
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        if current_length + len(sentence) > chunk_size:
            if current_chunk:
                chunks.append("。".join(current_chunk) + "。")
                current_chunk = []
                current_length = 0
                
        current_chunk.append(sentence)
        current_length += len(sentence)
        
    if current_chunk:
        chunks.append("。".join(current_chunk) + "。")
        
    return chunks

if __name__ == "__main__":
    analyze_document(r"C:\Users\Administrator\Desktop\工作区\路演文稿.docx")
    analyze_document(r"C:\Users\Administrator\Desktop\工作区\问题.docx") 